import streamlit as st
import pandas as pd
from docx import Document
import re
import os
import tempfile
import base64
import zipfile
import traceback
import datetime

# === Page Config ===
st.set_page_config(page_title="Employee Agreement Generator", layout="centered")


# === Base64 Logo ===
def get_base64_image(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()


# def convert_to_pdf(input_path, output_path):
#     output = pypandoc.convert_file(input_path, 'pdf', outputfile=output_path)
#     return output

# === Page Config ===
# st.set_page_config(page_title="Employee Agreement Generator", layout="centered")

# === Logo ===
img_base64 = get_base64_image(
    r"C:\Users\singhals\OneDrive - Zinnia\Automation\DocuWave\Zinnia_Logo_HORIZONTAL_Full_Color_RGB.png"
)

# === Simple HR Login ===
USERNAME = "HR001"
PASSWORD = "zinnia@2025"

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.title("🔐 HR Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    login_btn = st.button("Login")

    if login_btn:
        if username == USERNAME and password == PASSWORD:
            st.session_state.authenticated = True
            st.success("✅ Login successful!")
            st.rerun()
        else:
            st.error("❌ Invalid credentials. Please try again.")
    st.stop()

# if st.button("🚪 Logout", key="logout", help="Click to log out"):
#     st.session_state.authenticated = False
#     st.rerun()

if st.session_state.authenticated:
    st.markdown(
        """
        <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 30px;">
            <img src="data:image/png;base64,{img_base64}" width="140">
            <form action="?logout=true">
                <button style="background: #6c757d; color:white; border:none; padding:8px 16px; border-radius:5px; cursor:pointer;">
                    Logout
                </button>
            </form>
        </div>
    """.format(
            img_base64=img_base64
        ),
        unsafe_allow_html=True,
    )

# === App Title ===
st.title("HR Letters - DocuWave")


# === Load External CSS ===
def local_css(file_name):
    try:
        with open(file_name) as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    except FileNotFoundError:
        st.error(f"CSS file not found at: {file_name}")


local_css("style.css")

# === File Upload ===
uploaded_excel = st.file_uploader("📁 Upload Excel File", type=["xlsx"])
sheet_name = None
df = None
if uploaded_excel:
    excel_file = pd.ExcelFile(uploaded_excel)
    sheet_name = st.selectbox("📄 Select a Sheet", options=excel_file.sheet_names)
    if sheet_name:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        df.columns = [col.strip().lower().replace(" ", "_") for col in df.columns]

if df is not None and sheet_name:
    uploaded_template = st.file_uploader("📄 Upload Word Template", type=["docx"])
    if uploaded_template:

        def extract_placeholders(doc):
            text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            return set(re.findall(r"«(.*?)»", text))

        try:
            doc_check = Document(uploaded_template)
            placeholders_check = extract_placeholders(doc_check)
            required_columns = [
                ph.lower().replace(" ", "_").replace("-", "_")
                for ph in placeholders_check
            ]
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                st.error(
                    f"❌ The following required column(s) are missing in the Excel file: {', '.join(missing_columns)}"
                )
                st.stop()
            id_column = next((col for col in df.columns if "id" in col.lower()), None)
            name_column = next(
                (col for col in df.columns if "name" in col.lower()), None
            )
            if not id_column or not name_column:
                st.error("❌ Could not find 'ID' or 'Name' column in the Excel file.")
                st.stop()
            df["display_label"] = (
                df[name_column].astype(str).str.strip()
                + " ("
                + df[id_column].astype(str).str.strip()
                + ")"
            )
            display_to_id = dict(zip(df["display_label"], df[id_column]))
            df_with_sno = df.reset_index()
            df_with_sno["sno"] = df_with_sno.index + 1
            st.subheader("📋 Generation Mode")
            selected_mode = st.radio(
                "Choose how to generate agreements:",
                [
                    "Select Individual Employee",
                    "Select by Serial Number Range",
                    "All Employees",
                ],
            )
            if selected_mode == "Select Individual Employee":
                employee_options = df["display_label"].tolist()
                selected_display = st.selectbox(
                    "🧑‍💼 Select Employee", employee_options
                )
            elif selected_mode == "Select by Serial Number Range":
                col1, col2 = st.columns(2)
                with col1:
                    start_sno = st.number_input(
                        "From Row No.", min_value=1, max_value=len(df), value=1
                    )
                with col2:
                    end_sno = st.number_input(
                        "To Row No.",
                        min_value=start_sno,
                        max_value=len(df),
                        value=min(start_sno + 4, len(df)),
                    )
                st.info(
                    f"📊 Will generate agreements for {end_sno - start_sno + 1} employees (Row No. {start_sno} to {end_sno})"
                )
            generate_btn = st.button("📝 Generate Agreement(s)")

            def replace_text_in_paragraphs(paragraphs, replacements):
                for para in paragraphs:
                    full_text = "".join(run.text for run in para.runs)
                    for placeholder, value in replacements.items():
                        if placeholder in full_text:
                            if any(placeholder in run.text for run in para.runs):
                                for run in para.runs:
                                    run.text = run.text.replace(placeholder, value)
                            else:
                                replaced_text = full_text.replace(placeholder, value)
                                for run in para.runs:
                                    run.text = ""
                                para.runs[0].text = replaced_text

            def replace_text_in_tables(tables, replacements):
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_text_in_paragraphs(cell.paragraphs, replacements)

            def generate_document_for_employee(emp, template_doc):
                """Generate document for a single employee"""
                doc = Document(uploaded_template)
                placeholders = extract_placeholders(doc)

                replacements = {}
                for ph in placeholders:
                    norm_ph = ph.lower().replace(" ", "_").replace("-", "_")
                    match = next(
                        (col for col in df.columns if norm_ph == col or norm_ph in col),
                        None,
                    )

                    if match:
                        value = emp[match]
                        if pd.isna(value) or str(value).strip() == "":
                            value = ""  # Replace missing value with blank
                        elif isinstance(
                            value, (pd.Timestamp, datetime.datetime, datetime.date)
                        ):
                            try:
                                value = value.strftime(
                                    "%B %#d, %Y"
                                )  # e.g., July 1, 2025 (Windows compatible)
                            except Exception:
                                value = str(value)
                        replacements[f"«{ph}»"] = str(value)
                    else:
                        replacements[f"«{ph}»"] = ""  # If column not found, also blank

                replace_text_in_paragraphs(doc.paragraphs, replacements)
                replace_text_in_tables(doc.tables, replacements)

                emp_id = str(emp[id_column]).strip()
                emp_name = str(emp[name_column]).strip().title()
                file_name = f"{emp_id}_{emp_name}_{sheet_name}.docx"

                return doc, file_name

        except Exception as e:
            st.error(f"❌ Error: {e}")

        if generate_btn:
            if selected_mode == "All Employees":
                with st.spinner("Generating agreements for all employees..."):
                    with tempfile.TemporaryDirectory() as tmpdir:
                        output_files = []
                        for _, emp in df.iterrows():
                            doc, file_name = generate_document_for_employee(
                                emp, uploaded_template
                            )
                            doc_path = os.path.join(tmpdir, file_name)
                            doc.save(doc_path)
                            output_files.append(doc_path)

                        zip_path = os.path.join(tmpdir, "All_Agreements.zip")
                        with zipfile.ZipFile(zip_path, "w") as zipf:
                            for file in output_files:
                                zipf.write(file, arcname=os.path.basename(file))

                        with open(zip_path, "rb") as f:
                            st.success("✅ All agreements generated successfully.")
                            st.download_button(
                                label="⬇️ Download All Agreements (ZIP)",
                                data=f,
                                file_name="All_Agreements.zip",
                                mime="application/zip",
                            )

            elif selected_mode == "Select by Serial Number Range":
                with st.spinner(
                    f"Generating agreements for employees Row.No {start_sno} to {end_sno}..."
                ):
                    with tempfile.TemporaryDirectory() as tmpdir:
                        output_files = []
                        selected_employees = df.iloc[
                            start_sno - 1 : end_sno
                        ]  # Convert to 0-based indexing
                        for _, emp in selected_employees.iterrows():
                            doc, file_name = generate_document_for_employee(
                                emp, uploaded_template
                            )
                            doc_path = os.path.join(tmpdir, file_name)
                            doc.save(doc_path)
                            output_files.append(doc_path)
                        zip_path = os.path.join(
                            tmpdir, f"Agreements_SNO_{start_sno}_to_{end_sno}.zip"
                        )
                        with zipfile.ZipFile(zip_path, "w") as zipf:
                            for file in output_files:
                                zipf.write(file, arcname=os.path.basename(file))
                        with open(zip_path, "rb") as f:
                            st.success(
                                f"✅ Agreements generated for Row.No {start_sno} to {end_sno} successfully."
                            )
                            st.download_button(
                                label=f"⬇️ Download Agreements Row.No {start_sno}-{end_sno} (ZIP)",
                                data=f,
                                file_name=f"Agreements_SNO_{start_sno}_to_{end_sno}.zip",
                                mime="application/zip",
                            )

            else:  # Individual Employee
                selected_id = display_to_id[selected_display]
                emp = df[df[id_column] == selected_id].iloc[0]
                with st.spinner("Generating agreement for selected employee..."):
                    with tempfile.TemporaryDirectory() as tmpdir:
                        doc, file_name = generate_document_for_employee(
                            emp, uploaded_template
                        )
                        doc_path = os.path.join(tmpdir, file_name)
                        doc.save(doc_path)
                        with open(doc_path, "rb") as f:
                            st.success("✅ Agreement generated successfully.")
                            st.download_button(
                                label="⬇️ Download Agreement",
                                data=f,
                                file_name=file_name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )

else:
    st.info("📌 Please upload both an Excel file and a Word (.docx) template to begin.")
